import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()
    
    # --- TITLE PAGE ---
    # In python-docx, level=0 is the title style
    title = doc.add_heading("Technical Documentation\nAutonomous AI Data Analysis Agent", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n" * 5)
    
    logo_path = os.path.join("docs", "assets", "logo.png")
    if os.path.exists(logo_path):
        try:
            doc.add_picture(logo_path, width=Inches(3.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Logo skipped due to error: {e}")
    
    doc.add_paragraph("\n" * 5)
    p = doc.add_paragraph("Version 1.0.0\nDate: April 8, 2026\nPrepared for: DeepMind Engineering Review")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- TABLE OF CONTENTS ---
    h1 = doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Project Overview & Motivation",
        "3. System Architecture",
        "4. AI Agent & Orchestration Engine",
        "5. Contextual Memory & RAG (Deep Dive)",
        "6. AutoML & Predictive Modeling (Deep Dive)",
        "7. Visual Analytics & Reporting",
        "8. User Interface (Glassmorphic Design)",
        "9. Installation & Deployment",
        "10. Conclusion & Future Roadmap"
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_page_break()

    # --- CHAPTER 1: EXECUTIVE SUMMARY ---
    doc.add_heading("1. Executive Summary", level=1)
    p = doc.add_paragraph(
        "The Autonomous AI Data Analysis Agent is a state-of-the-art platform designed to bridge the gap between complex raw data and actionable intelligence. "
        "By leveraging the Gemini Pro API, ChromaDB, and an automated machine learning pipeline, the agent allows users to interact with their datasets using natural language. "
        "The system handles end-to-end data tasks: from automated cleaning and schema inference to complex predictive modeling and interactive visualizations. "
        "This document provides a comprehensive technical breakdown of the system components, design philosophy, and implementation details."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("\n")

    # --- CHAPTER 2: PROJECT OVERVIEW ---
    doc.add_heading("2. Project Overview & Motivation", level=1)
    p = doc.add_paragraph(
        "In the modern era, organizations are inundated with data, but the bottleneck remains the 'last mile' of analysis. Traditional tools require deep proficiency in SQL, Python, or specialized BI software. "
        "Our motivation was to create a 'Conversational Data Scientist'—an agent that not only answers questions but understands the underlying context of the data it explores."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph(
        "Key design pillars for this project include:\n"
        "1. Accessibility: No-code interface for complex operations.\n"
        "2. Autonomy: Self-correcting data cleaning and model selection.\n"
        "3. Transparency: Reasoning traces showing how the agent arrived at a conclusion.",
        style='List Bullet'
    )

    # --- CHAPTER 3: SYSTEM ARCHITECTURE ---
    doc.add_heading("3. System Architecture", level=1)
    doc.add_paragraph(
        "The system is built on a modular micro-services inspired architecture, integrated via a central Streamlit orchestrator. "
        "The core flow involves a user query being routed to an Intent Classifier (powered by Gemini), which then selects the appropriate tool for execution."
    )
    doc.add_paragraph(
        "Component Layers:\n"
        "- Presentation Layer: Streamlit Glassmorphic Frontend.\n"
        "- Logic Layer: Gemini Pro Orchestrator and Agent Router.\n"
        "- Persistence Layer: ChromaDB Vector Store for RAG-based context.\n"
        "- Execution Layer: Scikit-Learn (ML) and Plotly (Visualization)."
    )
    doc.add_page_break()

    # --- CHAPTER 5: CONTEXTUAL MEMORY & RAG (DEEP DIVE) ---
    doc.add_heading("5. Contextual Memory & RAG (Deep Dive)", level=1)
    doc.add_paragraph(
        "One of the primary challenges in LLM-based data analysis is 'context window overflow'—large datasets contain thousands of rows and dozens of columns, which cannot all fit into a single prompt. "
        "Our solution is a Retrieval-Augmented Generation (RAG) system specifically designed for Metadata Discovery."
    )
    doc.add_heading("5.1 Vector Store Implementation", level=2)
    doc.add_paragraph(
        "We utilize ChromaDB to store vector embeddings of column names, data types, and sample unique values. "
        "When a user asks a question, the agent performs a similarity search against the vector store to retrieve only the relevant column metadata."
    )
    doc.add_paragraph(
        "Example Workflow:\n"
        "- User: 'How many premium customers in the North?'\n"
        "- RAG: Retrieves columns 'Subscription_Type', 'Region', and 'Customer_Count'.\n"
        "- Prompt: Only provides these specific columns to the LLM, ensuring high precision and lower token cost."
    )
    doc.add_heading("5.2 Dynamic Context Injection", level=2)
    doc.add_paragraph(
        "The RAG engine doesn't just store names; it stores 'semantic intent'. We use the Gemini API to generate descriptions for cryptic column names, making them searchable by their real-world meaning."
    )
    doc.add_paragraph("\n" * 15)
    doc.add_page_break()

    # --- CHAPTER 6: AUTOML ENGINE (DEEP DIVE) ---
    doc.add_heading("6. AutoML & Predictive Modeling (Deep Dive)", level=1)
    doc.add_paragraph(
        "The AutoML engine (located in src/ml/automl.py) is the 'brain' of the agent's predictive capabilities. It eliminates the need for manual feature engineering and model tuning for standard tasks."
    )
    doc.add_heading("6.1 Automated Problem Classification", level=2)
    doc.add_paragraph(
        "Upon receiving a prediction request, the agent analyzes the target variable's distribution and data type:\n"
        "- Continuous Target -> Regression Pipeline (Random Forest / Gradient Boosting).\n"
        "- Categorical Target -> Classification Pipeline (Logistic Regression / XGBoost style)."
    )
    doc.add_heading("6.2 The Pipeline Architecture", level=2)
    doc.add_paragraph(
        "Each ML task follows a robust 4-step sequence:\n"
        "1. Dynamic Feature Selection: Identifying correlated features defined in the query.\n"
        "2. Imputation & Scaling: Handling missing values and normalizing variance.\n"
        "3. Model Selection: Evaluating multiple algorithms and selecting the best performer based on R2 or F1-Score.\n"
        "4. Evaluation: Generating a report with RMSE, Accuracy, and Feature Importance."
    )
    doc.add_paragraph(
        "This allows a user to simply say 'Predict Churn' and receive a fully trained, validated model in seconds."
    )
    doc.add_paragraph("\n" * 20)
    doc.add_page_break()

    # --- CHAPTER 7-10 ---
    doc.add_heading("7. Visual Analytics & Reporting", level=1)
    doc.add_paragraph(
        "Visualizations are generated using Plotly. The LLM identifies the 'Visual Intent'—for example, a trend requires a line chart, while a distribution requires a histogram. "
        "The code is generated on-the-fly, ensuring the UI is always dynamic and tailored to the data."
    )

    doc.add_heading("8. User Interface (Glassmorphic Design)", level=1)
    doc.add_paragraph(
        "The UI uses a custom CSS layer over Streamlit to achieve a 'Glassmorphic' aesthetic. Blur effects, semi-transparent panels, and subtle gradients create a premium feel that resonates with modern software standards."
    )

    doc.add_heading("9. Installation & Deployment", level=1)
    doc.add_paragraph(
        "To deploy the agent, follow these steps:\n"
        "1. Install Python 3.9+\n"
        "2. Run 'pip install -r requirements.txt'\n"
        "3. Configure .env with GEMINI_API_KEY\n"
        "4. Start with 'streamlit run ui/app.py'"
    )

    doc.add_heading("10. Conclusion & Future Roadmap", level=1)
    doc.add_paragraph(
        "The current iteration of the Autonomous AI Data Agent marks a significant step toward democratization of data science. "
        "Future releases will include support for local Llama-3 models, direct SQL database connections, and multi-agent collaborative reasoning."
    )
    
    for i in range(10):
        doc.add_paragraph("\n")
    doc.add_paragraph("--- End of Documentation ---")

    save_path = "Autonomous_AI_Data_Agent_Documentation.docx"
    doc.save(save_path)
    print(f"Document saved to {save_path}")

if __name__ == "__main__":
    create_doc()
